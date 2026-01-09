from docx import Document
from docx.shared import Inches, Pt, RGBColor
import os

def generate_word_doc(paper, questions, output_path):
    document = Document()
    
    # Title
    heading = document.add_heading(f'Solutions: {paper.filename}', 0)
    heading.alignment = 1 # Center
    
    for idx, q in enumerate(questions):
        # Question Header
        p_head = document.add_heading(f'Question {idx + 1}', level=1)
        
        # Check integrity
        if getattr(q, 'is_incomplete', False):
             p_warn = document.add_paragraph("⚠️ This question was marked as incomplete/cut-off.")
             p_warn.runs[0].font.color.rgb = RGBColor(255, 165, 0) # Orange
        
        # 1. Question Text
        if q.ocr_text:
            document.add_heading('Problem Statement', level=2)
            p_text = document.add_paragraph(q.ocr_text)
            p_text.paragraph_format.left_indent = Inches(0.25)
            
        # 2. Image (Only if it's a specific crop, not the full paper)
        # We assume crops have 'crop_' in filename or are different from paper.file_path
        # But here checking filename is safer if we stick to convention
        # Or simple check: if q.image_path != paper.file_path
        if q.image_path and q.image_path != paper.file_path:
            abs_image_path = os.path.abspath(q.image_path)
            if os.path.exists(abs_image_path):
                try:
                    document.add_picture(abs_image_path, width=Inches(3.0))
                except Exception:
                    pass
        
        # 3. Answer
        answer_text = getattr(q, 'answer', '')
        if answer_text:
            document.add_heading('Answer', level=2)
            p_ans = document.add_paragraph()
            run_ans = p_ans.add_run(answer_text)
            run_ans.bold = True
            run_ans.font.size = Pt(12)
            run_ans.font.color.rgb = RGBColor(0, 100, 0) # Dark Green

        # 4. Analysis
        analysis_text = getattr(q, 'analysis', '')
        # Fallback to solution_text if analysis is empty
        if not analysis_text:
            analysis_text = q.solution_text
            
        if analysis_text:
            document.add_heading('Analysis', level=2)
            document.add_paragraph(analysis_text)
        elif not getattr(q, 'is_incomplete', False):
            document.add_paragraph("[Analysis generating or not available]", style='Italic')
            
        document.add_page_break()
        
    document.save(output_path)
    return output_path
